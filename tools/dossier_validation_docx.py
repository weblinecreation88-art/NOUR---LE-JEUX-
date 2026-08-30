# -*- coding: utf-8 -*-
"""
Genere VALIDATION_RELIGIEUSE.docx — version Word annotable du dossier.

Meme contenu que le .md, mais dans un format ou un relecteur peut ecrire
directement (zones de remarques, cases a cocher).

Usage : python tools/dossier_validation_docx.py
"""
import io
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

RACINE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

OCRE = RGBColor(0xB0, 0x7C, 0x2A)
GRIS = RGBColor(0x6B, 0x57, 0x44)
ROUGE = RGBColor(0xB5, 0x62, 0x3C)
VERT = RGBColor(0x5A, 0x7A, 0x3B)


def lire(nom):
    with io.open(os.path.join(RACINE, 'data', nom + '.json'),
                 encoding='utf-8') as f:
        return json.load(f)


def para(doc, texte='', taille=11, gras=False, italique=False,
         couleur=None, avant=0, apres=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(avant)
    p.paragraph_format.space_after = Pt(apres)
    r = p.add_run(texte)
    r.font.size = Pt(taille)
    r.bold = gras
    r.italic = italique
    if couleur is not None:
        r.font.color.rgb = couleur
    return p


def champ_remarques(doc):
    """Zone encadree ou le relecteur ecrit."""
    para(doc, 'Remarques du relecteur :', 9, gras=True, couleur=GRIS,
         avant=6, apres=2)
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    cell.text = ''
    # Trois lignes vides pour ecrire
    for _ in range(3):
        cell.add_paragraph()
    para(doc, 'Statut :   [  ] VALIDÉ        [  ] À CORRIGER        '
              '[  ] À RETIRER', 10, gras=True, avant=4, apres=12)


def tableau_meta(doc, lignes):
    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'
    for cle, val in lignes:
        row = t.add_row().cells
        row[0].text = cle
        row[1].text = str(val)
        for p in row[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        for p in row[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
    doc.add_paragraph()
    return t


def main():
    quiz = lire('quiz')
    know = lire('knowledge')
    scenes = lire('scenes')

    emplacement = {}
    for sc in scenes['scenes']:
        for b in sc['beats']:
            if b.get('type') == 'quiz':
                emplacement[b['question']] = '%d. %s' % (sc['ordre'],
                                                         sc['titre'])

    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(11)

    # ------------------------------------------------------------ couverture
    t = para(doc, 'NOUR — Le Jeu', 26, gras=True, couleur=OCRE, apres=2)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = para(doc, 'Dossier de validation du contenu religieux', 15,
             couleur=GRIS, apres=2)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = para(doc, 'Chapitre 1 — La Lumière perdue', 11, italique=True,
             couleur=GRIS, apres=24)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc, 'À l’attention de la personne chargée de la relecture',
         13, gras=True, apres=8)
    para(doc,
         'Ce document rassemble la totalité du contenu religieux présenté '
         'au joueur dans le Chapitre 1. Rien d’autre n’est affiché.', 11)
    para(doc,
         'Il est extrait automatiquement des fichiers du jeu : le texte '
         'ci-dessous est exactement celui qui s’affiche à l’écran.', 11,
         apres=16)

    para(doc, 'Ce qui est demandé', 13, gras=True, couleur=OCRE, apres=6)
    for i, txt in enumerate([
        'La formulation est-elle correcte ? (sinon, proposer la correction)',
        'La référence est-elle exacte ? (sourate/verset, recueil/numéro)',
        'Le contenu peut-il être présenté ainsi à un jeune public ?',
        'Indiquer le statut : VALIDÉ, À CORRIGER ou À RETIRER.',
    ], 1):
        para(doc, '%d.  %s' % (i, txt), 11, apres=3)

    para(doc, '', apres=10)
    para(doc, 'Règles déjà appliquées dans le jeu', 13, gras=True,
         couleur=OCRE, apres=6)
    para(doc, 'Elles sont vérifiées automatiquement par les tests du projet :',
         10, italique=True, couleur=GRIS, apres=6)
    for txt in [
        'Aucun hadith ni verset inventé.',
        'Aucune parole attribuée au Prophète ﷺ sans référence.',
        'Aucune invocation présentée comme une attaque, un sort ou un pouvoir.',
        'Aucune représentation de prophète.',
        'L’XP est une mécanique de jeu : elle ne mesure jamais la foi.',
        'Une mauvaise réponse ne retire rien et ne culpabilise pas.',
        'Aucune preuve d’acte religieux n’est demandée '
        '(ni photo, ni audio, ni localisation).',
    ]:
        para(doc, '•  ' + txt, 10, apres=3)

    para(doc, '', apres=8)
    para(doc,
         'Tout contenu non validé est affiché au joueur avec la mention '
         '« Référence à valider ». Rien n’est présenté comme définitif.',
         11, gras=True, couleur=ROUGE)

    # ------------------------------------------------------------------ quiz
    doc.add_page_break()
    para(doc, 'Partie 1 — Les questions de quiz', 18, gras=True,
         couleur=OCRE, apres=4)
    para(doc, '%d questions. La bonne réponse est marquée d’une coche.'
         % len(quiz['questions']), 10, italique=True, couleur=GRIS, apres=14)

    for i, q in enumerate(quiz['questions'], 1):
        statut = 'À VALIDER' if q['statut'] == 'A_VALIDER' else 'VALIDE'
        para(doc, 'Q%d — %s' % (i, q['id']), 14, gras=True, avant=10, apres=6)
        tableau_meta(doc, [
            ('Notion', q['notion']),
            ('Scène', emplacement.get(q['id'], '—')),
            ('Source indiquée', q['source']),
            ('Référence indiquée', q['reference']),
            ('Statut actuel', statut),
        ])

        para(doc, 'Question posée au joueur', 10, gras=True, couleur=GRIS,
             apres=2)
        para(doc, q['question'], 12, apres=8)

        para(doc, 'Réponses proposées', 10, gras=True, couleur=GRIS, apres=2)
        for j, c in enumerate(q['choix']):
            bon = (j == q['reponse'])
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run('%s.  %s' % (chr(65 + j), c))
            r.font.size = Pt(11)
            if bon:
                r.bold = True
                r.font.color.rgb = VERT
                rr = p.add_run('    ✔ bonne réponse')
                rr.font.size = Pt(9)
                rr.font.color.rgb = VERT

        para(doc, '', apres=6)
        para(doc, 'Explication affichée après la réponse', 10, gras=True,
             couleur=GRIS, apres=2)
        para(doc, q['explication'], 11, italique=True, apres=6)

        champ_remarques(doc)
        if i < len(quiz['questions']):
            doc.add_page_break()

    # --------------------------------------------------------------- notions
    doc.add_page_break()
    para(doc, 'Partie 2 — Les fiches de notions', 18, gras=True,
         couleur=OCRE, apres=4)
    para(doc, '%d fiches, consultables par le joueur dans la bibliothèque '
              'du jeu.' % len(know['notions']), 10, italique=True,
         couleur=GRIS, apres=14)

    for i, n in enumerate(know['notions'], 1):
        statut = 'À VALIDER' if n['statut'] == 'A_VALIDER' else 'VALIDE'
        para(doc, 'N%d — %s   %s' % (i, n['translitteration'], n['terme_ar']),
             14, gras=True, avant=10, apres=6)
        tableau_meta(doc, [
            ('Identifiant', n['id']),
            ('Titre français', n['titre_fr']),
            ('Niveau', n['niveau']),
            ('Source indiquée', n['source']),
            ('Référence indiquée', n['reference']),
            ('Statut actuel', statut),
        ])
        para(doc, 'Définition affichée', 10, gras=True, couleur=GRIS, apres=2)
        para(doc, n['definition'], 11, apres=8)
        para(doc, 'Explication affichée', 10, gras=True, couleur=GRIS, apres=2)
        para(doc, n['explication'], 11, italique=True, apres=6)
        champ_remarques(doc)
        if i < len(know['notions']):
            doc.add_page_break()

    # ---------------------------------------------------------------- points
    doc.add_page_break()
    para(doc, 'Partie 3 — Points signalés par l’équipe', 18, gras=True,
         couleur=OCRE, apres=6)
    para(doc, 'Ces entrées sont celles où nous avons le moins de certitude. '
              'Elles méritent une attention particulière.', 11, apres=10)

    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    entetes = t.rows[0].cells
    entetes[0].text = 'Entrée'
    entetes[1].text = 'Pourquoi'
    for c in entetes:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for cle, pourquoi in [
        ('q4_sabr / sabr',
         'Aucune source précise. Le cahier des charges indique '
         '« à vérifier précisément ».'),
        ('q_niyyah / niyyah',
         '« Référence à documenter » — le hadith habituellement cité '
         'n’a pas été rattaché formellement.'),
        ('q9_shukr / shukr', '« Référence à documenter ».'),
        ('q7_salam / salam',
         'Muwatta Malik cité, mais numéro non vérifié.'),
        ('q6_adab / adab', 'Bukhari 6136 cité, à confirmer.'),
        ('q1_taaruf, q3_istiadhah',
         'Références coraniques (49:13 ; 16:98) qui semblent solides, '
         'mais à confirmer formellement.'),
    ]:
        row = t.add_row().cells
        row[0].text = cle
        row[1].text = pourquoi
        for p in row[0].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.bold = True
        for p in row[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

    para(doc, '', apres=16)
    para(doc, 'Partie 4 — Contenu volontairement NON religieux', 15,
         gras=True, couleur=OCRE, apres=6)
    para(doc, 'Une seule entrée est marquée VALIDE dans le jeu : q10_ilm.',
         11, apres=4)
    para(doc, 'Ce n’est pas un texte religieux. C’est la phrase de '
              'philosophie du jeu, écrite par l’équipe :', 11, apres=4)
    para(doc, '« Le ʿilm, les efforts et le fait de passer à l’action. »',
         12, italique=True, gras=True, apres=6)
    para(doc, 'Elle est présentée comme telle au joueur, jamais comme une '
              'citation. Merci de confirmer que cette distinction est '
              'claire.', 11, apres=16)

    para(doc, 'Après la relecture', 15, gras=True, couleur=OCRE, apres=6)
    para(doc, 'Renvoyer ce document annoté. Les corrections seront '
              'reportées dans les fichiers du jeu, et le statut passera à '
              'VALIDE uniquement pour les entrées explicitement validées.',
         11, apres=4)
    para(doc, 'Tant qu’une entrée reste à valider, le jeu continue '
              'd’afficher « Référence à valider » au joueur.', 11,
         gras=True, couleur=ROUGE)

    dest = os.path.join(RACINE, 'VALIDATION_RELIGIEUSE.docx')
    doc.save(dest)
    print('VALIDATION_RELIGIEUSE.docx : %.0f Ko'
          % (os.path.getsize(dest) / 1024))
    print('  %d questions + %d notions'
          % (len(quiz['questions']), len(know['notions'])))


if __name__ == '__main__':
    main()
